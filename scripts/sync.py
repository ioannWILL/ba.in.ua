#!/usr/bin/env python3
"""
JIRA → WordPress translation workflow.

Daily checks for JIRA issues in 'To Review' status, then for each one:
  - Downloads .docx translation attachment
  - Extracts Ukrainian title + body
  - Uploads featured image to WordPress
  - Generates bilingual tags via Claude
  - Creates WordPress draft post with attribution footer
  - Comments on JIRA issue with WordPress draft link
  - Transitions JIRA issue to 'In Review'
  - Translator is read from the Assignee field
"""

import json
import logging
import mimetypes
import os
import re
from io import BytesIO

import anthropic
import requests
from bs4 import BeautifulSoup
from docx import Document

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
log = logging.getLogger(__name__)

# ── Configuration ──────────────────────────────────────────────────────────────

JIRA_BASE_URL         = os.environ["JIRA_BASE_URL"]          # e.g. https://yourteam.atlassian.net
JIRA_EMAIL            = os.environ["JIRA_EMAIL"]
JIRA_API_TOKEN        = os.environ["JIRA_API_TOKEN"]
JIRA_PROJECT_KEY      = os.environ["JIRA_PROJECT_KEY"]
JIRA_TRIGGER_STATUS   = os.environ.get("JIRA_TRIGGER_STATUS") or "TO REVIEW"
JIRA_NEXT_STATUS      = os.environ.get("JIRA_NEXT_STATUS") or "IN REVIEW"
# Custom field ID for "Translated by" — find it via:
# GET /rest/api/3/field  (look for your custom field name)
JIRA_TRANSLATOR_FIELD = os.environ.get("JIRA_TRANSLATOR_FIELD", "")
# Optional custom field ID for the original article URL
JIRA_ORIGINAL_URL_FIELD = os.environ.get("JIRA_ORIGINAL_URL_FIELD", "")

WP_BASE_URL    = os.environ["WP_BASE_URL"]          # e.g. https://ba.in.ua
WP_USERNAME    = os.environ["WP_USERNAME"]
WP_APP_PASSWORD = os.environ["WP_APP_PASSWORD"]

ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]

REVIEWER_NAME = "Іван Вільчавський"
REVIEWER_URL  = "https://www.linkedin.com/in/ivan-vilchavskyi"


# ── JIRA client ────────────────────────────────────────────────────────────────

class JiraClient:
    def __init__(self):
        self.session = requests.Session()
        self.session.auth = (JIRA_EMAIL, JIRA_API_TOKEN)
        self.session.headers.update({"Accept": "application/json"})
        self.base = JIRA_BASE_URL.rstrip("/")
        self._field_ids = self._discover_field_ids()

    def _discover_field_ids(self) -> dict[str, str]:
        """Build name→id map for all JIRA fields (used to resolve custom fields by name)."""
        try:
            fields = self._get("/rest/api/3/field")
            mapping = {f["name"].lower(): f["id"] for f in fields}
            log.info("Discovered %d JIRA fields", len(mapping))
            return mapping
        except Exception as e:
            log.warning("Could not discover JIRA fields: %s", e)
            return {}

    def _field_id(self, name: str) -> str:
        """Return field ID for a given field name (case-insensitive)."""
        return self._field_ids.get(name.lower(), "")

    def _get(self, path, **kw):
        r = self.session.get(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r.json()

    def _post(self, path, **kw):
        r = self.session.post(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r

    def get_issues_in_status(self, status: str) -> list[dict]:
        jql = f'project = "{JIRA_PROJECT_KEY}" AND status = "{status}" ORDER BY updated ASC'
        url = f"{self.base}/rest/api/3/search/jql"
        # Request standard fields + any known custom fields
        custom = [f for f in [
            self._field_id("link to original source"),
            self._field_id("translated by"),
        ] if f]
        fields = ["summary", "status", "assignee", "attachment", "description"] + custom
        log.info("JIRA search: jql=%s  extra_fields=%s", jql, custom)
        r = self.session.post(url, json={"jql": jql, "maxResults": 50, "fields": fields})
        log.info("JIRA search response: status=%s", r.status_code)
        r.raise_for_status()
        return r.json().get("issues", [])

    def get_attachments(self, issue: dict) -> list[dict]:
        return issue.get("fields", {}).get("attachment", [])

    def download_attachment(self, attachment: dict) -> bytes:
        r = self.session.get(attachment["content"])
        r.raise_for_status()
        return r.content

    def get_translator(self, issue: dict) -> str:
        """Return translator name from the custom 'Translated by' field."""
        fields = issue.get("fields", {})
        # Try env var override first, then auto-discovered field ID
        translator_field = JIRA_TRANSLATOR_FIELD or self._field_id("translated by")
        if translator_field:
            val = fields.get(translator_field)
            # Multi-user picker returns a list
            if isinstance(val, list) and val:
                return val[0].get("displayName") or val[0].get("name") or ""
            if isinstance(val, dict):
                return val.get("displayName") or val.get("name") or ""
            if isinstance(val, str):
                return val
        # Fallback: assignee
        user = fields.get("assignee")
        if isinstance(user, dict):
            return user.get("displayName") or user.get("name") or ""
        return ""

    def get_original_url(self, issue: dict) -> str:
        """Extract the original English article URL from the JIRA issue."""
        fields = issue.get("fields", {})
        issue_key = issue.get("key", "")

        # 1. Custom field "Link to original source" (env var or auto-discovered)
        url_field = JIRA_ORIGINAL_URL_FIELD or self._field_id("link to original source")
        if url_field:
            val = fields.get(url_field)
            if isinstance(val, str) and val.startswith("http"):
                log.info("Original URL from custom field: %s", val)
                return val

        # 2. JIRA remote links (web links added via "Link" button)
        if issue_key:
            try:
                links = self._get(f"/rest/api/3/issue/{issue_key}/remotelink")
                for link in links:
                    url = link.get("object", {}).get("url", "")
                    if url and "atlassian.net" not in url and not _is_media_url(url):
                        log.info("Found remote link: %s", url)
                        return url
            except Exception as e:
                log.warning("Could not fetch remote links for %s: %s", issue_key, e)

        # 3. ADF hyperlink marks in description — prefer article URLs over image URLs
        desc = fields.get("description") or ""
        if isinstance(desc, dict):
            adf_urls = _adf_extract_urls(desc)
            article_urls = [
                u for u in adf_urls
                if "atlassian.net" not in u and not _is_media_url(u)
            ]
            if article_urls:
                log.info("Found article URL in ADF description: %s", article_urls[0])
                return article_urls[0]
            desc = _adf_to_text(desc)

        # 4. Plain-text URL in description
        urls = re.findall(r"https?://[^\s\]\)>\"'<]+", str(desc))
        article_urls = [u for u in urls if "atlassian.net" not in u and not _is_media_url(u)]
        if article_urls:
            return article_urls[0]
        return ""

    def get_transition_id(self, issue_key: str, status_name: str) -> str | None:
        data = self._get(f"/rest/api/3/issue/{issue_key}/transitions")
        for t in data.get("transitions", []):
            if t["to"]["name"].lower() == status_name.lower():
                return t["id"]
        return None

    def transition_issue(self, issue_key: str, status_name: str) -> bool:
        tid = self.get_transition_id(issue_key, status_name)
        if not tid:
            log.warning("Transition '%s' not found for %s", status_name, issue_key)
            return False
        self._post(
            f"/rest/api/3/issue/{issue_key}/transitions",
            json={"transition": {"id": tid}},
        )
        log.info("Transitioned %s → %s", issue_key, status_name)
        return True

    def add_comment(self, issue_key: str, text: str) -> None:
        self._post(
            f"/rest/api/3/issue/{issue_key}/comment",
            json={"body": {"type": "doc", "version": 1,
                           "content": [{"type": "paragraph",
                                        "content": [{"type": "text", "text": text}]}]}},
        )


_MEDIA_HOSTS = {
    "freepik.com", "unsplash.com", "shutterstock.com", "pixabay.com",
    "istockphoto.com", "gettyimages.com", "pexels.com", "flickr.com",
    "imgur.com", "cloudinary.com", "depositphotos.com", "dreamstime.com",
}
_MEDIA_EXT_RE = re.compile(r"\.(jpe?g|png|gif|webp|svg|bmp|tiff?)(\?|$)", re.I)


def _is_media_url(url: str) -> bool:
    """Return True if the URL points to an image hosting site or image file."""
    from urllib.parse import urlparse
    parsed = urlparse(url)
    host = parsed.netloc.lower().lstrip("www.")
    if any(host == h or host.endswith("." + h) for h in _MEDIA_HOSTS):
        return True
    if _MEDIA_EXT_RE.search(parsed.path):
        return True
    return False


def _adf_to_text(node: dict | list) -> str:
    """Recursively extract plain text from Atlassian Document Format."""
    if isinstance(node, list):
        return " ".join(_adf_to_text(n) for n in node)
    if isinstance(node, dict):
        if node.get("type") == "text":
            return node.get("text", "")
        return " ".join(_adf_to_text(c) for c in node.get("content", []))
    return ""


def _adf_extract_urls(node: dict | list) -> list[str]:
    """Extract all href URLs from ADF link marks (hyperlinks in rich text)."""
    urls: list[str] = []
    if isinstance(node, list):
        for n in node:
            urls.extend(_adf_extract_urls(n))
    elif isinstance(node, dict):
        for mark in node.get("marks", []):
            if mark.get("type") == "link":
                href = mark.get("attrs", {}).get("href", "")
                if href:
                    urls.append(href)
        for child in node.get("content", []):
            urls.extend(_adf_extract_urls(child))
    return urls


# ── .docx → HTML ───────────────────────────────────────────────────────────────

def _is_separator(text: str) -> bool:
    """Return True if the line is a visual separator (dashes, underscores, etc.)."""
    return bool(re.fullmatch(r"[-–—_=*#\s]{3,}", text))


def _is_cyrillic_title(text: str) -> bool:
    """Return True if text is a plausible Ukrainian title (more Cyrillic than Latin)."""
    cyrillic = sum(1 for c in text if "\u0400" <= c <= "\u04FF")
    latin = sum(1 for c in text if c.isascii() and c.isalpha())
    return cyrillic >= 3 and cyrillic >= latin


def docx_to_html(data: bytes) -> tuple[str, str]:
    """Parse .docx bytes and return (title, body_html).

    Title priority:
      1. First paragraph with Heading 1 / Title style
      2. First predominantly-Cyrillic paragraph (skips English metadata lines)
      3. First non-empty, non-separator paragraph as last resort
    """
    doc = Document(BytesIO(data))
    title = ""
    body_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or _is_separator(text):
            continue
        style = para.style.name.lower()

        if not title and (style.startswith("heading 1") or style == "title"):
            title = text
            continue

        body_parts.append(_para_to_html(para))

    if not title:
        # Pass 1: first Cyrillic-dominant paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and not _is_separator(text) and _is_cyrillic_title(text):
                title = text
                first_html = _para_to_html(para)
                if first_html in body_parts:
                    body_parts.remove(first_html)
                break

    if not title:
        # Pass 2: absolute fallback — first non-empty paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and not _is_separator(text):
                title = text
                first_html = _para_to_html(para)
                if first_html in body_parts:
                    body_parts.remove(first_html)
                break

    log.info("Extracted title from docx: %s", title)
    return title, "\n".join(body_parts)


def _para_to_html(para) -> str:
    HEADING_MAP = {
        "heading 1": "h1", "heading 2": "h2", "heading 3": "h3",
        "heading 4": "h4", "heading 5": "h5", "heading 6": "h6",
    }
    style = para.style.name.lower()
    tag = HEADING_MAP.get(style, "p")
    inner = _runs_to_html(para.runs)
    if not inner.strip():
        return ""
    return f"<{tag}>{inner}</{tag}>"


def _runs_to_html(runs) -> str:
    parts = []
    for run in runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts)


# ── WordPress client ───────────────────────────────────────────────────────────

class WpClient:
    def __init__(self):
        self.session = requests.Session()
        self.session.auth = (WP_USERNAME, WP_APP_PASSWORD)
        self.base = WP_BASE_URL.rstrip("/") + "/wp-json/wp/v2"

    def _get(self, path, **kw):
        r = self.session.get(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r.json()

    def _post(self, path, **kw):
        r = self.session.post(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r.json()

    def upload_media(self, filename: str, data: bytes, mime: str) -> int:
        result = self._post(
            "/media",
            data=data,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Type": mime,
            },
        )
        return result["id"]

    def get_or_create_tag(self, name: str) -> int:
        existing = self._get("/tags", params={"search": name, "per_page": 10})
        for tag in existing:
            if tag["name"].lower() == name.lower():
                return tag["id"]
        result = self._post("/tags", json={"name": name})
        return result["id"]

    def create_post(
        self,
        title: str,
        content: str,
        status: str = "draft",
        featured_media: int | None = None,
        tags: list[int] | None = None,
    ) -> dict:
        payload: dict = {"title": title, "content": content, "status": status}
        if featured_media:
            payload["featured_media"] = featured_media
        if tags:
            payload["tags"] = tags
        return self._post("/posts", json=payload)


# ── Tag generation via Claude ──────────────────────────────────────────────────

def generate_tags(title: str, body_html: str) -> list[str]:
    """Ask Claude to generate 3-4 English + 3-4 Ukrainian tags."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    plain_body = BeautifulSoup(body_html, "html.parser").get_text()[:3000]
    prompt = (
        "You are tagging an article for a Ukrainian business/analytics website (ba.in.ua).\n\n"
        f"Article title: {title}\n"
        f"Article excerpt:\n{plain_body}\n\n"
        "Generate exactly 3-4 concise tags in English and 3-4 matching tags in Ukrainian.\n"
        "Return ONLY a JSON array of strings, e.g. "
        '["Strategy", "Leadership", "Growth", "Стратегія", "Лідерство", "Зростання"]\n'
        "Tags must be 1-3 words, relevant, and suitable for a blog taxonomy. No explanations."
    )
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=256,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    match = re.search(r"\[.*?\]", raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    log.warning("Could not parse tags from Claude response: %s", raw)
    return []


# ── Image selection ────────────────────────────────────────────────────────────

def pick_main_image(
    attachments: list[dict], original_url: str, jira: "JiraClient"
) -> tuple[bytes, str, str] | None:
    """Return (image_bytes, filename, mime_type) for the best main image."""
    images = [a for a in attachments if a.get("mimeType", "").startswith("image/")]
    if not images:
        return None
    if len(images) == 1:
        att = images[0]
        return jira.download_attachment(att), att["filename"], att["mimeType"]

    # Multiple images: try to match with original article's og:image
    if original_url:
        og_url = _fetch_og_image(original_url)
        if og_url:
            og_name = og_url.split("/")[-1].split("?")[0].lower()
            for att in images:
                if att["filename"].lower() == og_name:
                    return jira.download_attachment(att), att["filename"], att["mimeType"]
            # Download directly from original article
            try:
                r = requests.get(og_url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
                if r.ok and r.headers.get("content-type", "").startswith("image/"):
                    mime = r.headers["content-type"].split(";")[0].strip()
                    ext = mimetypes.guess_extension(mime) or ".jpg"
                    return r.content, f"featured{ext}", mime
            except Exception as e:
                log.warning("Could not download og:image: %s", e)

    # Fallback: largest attachment
    att = max(images, key=lambda a: a.get("size", 0))
    return jira.download_attachment(att), att["filename"], att["mimeType"]


def _fetch_og_image(url: str) -> str:
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        for meta in (
            soup.find("meta", property="og:image"),
            soup.find("meta", attrs={"name": "twitter:image"}),
        ):
            if meta and meta.get("content"):
                return meta["content"]
        # Last resort: first image in article body
        for img in soup.select("article img, .entry-content img, .post-content img"):
            src = img.get("src") or img.get("data-src")
            if src and not src.endswith(".svg"):
                return src
    except Exception as e:
        log.warning("Could not fetch og:image from %s: %s", url, e)
    return ""


def _fetch_page_title(url: str) -> str:
    if not url:
        return ""
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        og = soup.find("meta", property="og:title")
        if og and og.get("content"):
            return og["content"].strip()
        if soup.title and soup.title.string:
            return soup.title.string.strip()
    except Exception as e:
        log.warning("Could not fetch page title from %s: %s", url, e)
    return ""


# ── Attribution paragraph ──────────────────────────────────────────────────────

def build_attribution(original_url: str, original_title: str, translator: str) -> str:
    if original_url and original_title:
        article_part = f'<a href="{original_url}">{original_title}</a>'
    elif original_url:
        article_part = f'<a href="{original_url}">{original_url}</a>'
    else:
        article_part = original_title or "оригінальна стаття"

    reviewer_part = f'<a href="{REVIEWER_URL}">{REVIEWER_NAME}</a>'
    translator_part = translator or "невідомий перекладач"

    return (
        f"<p>Оригінальна стаття — {article_part}, "
        f"переклад — {translator_part}, "
        f"ревью — {reviewer_part}. "
        f"Зображення з оригінальної статті.</p>"
    )


# ── Main issue processor ───────────────────────────────────────────────────────

def process_issue(issue: dict, jira: JiraClient, wp: WpClient) -> bool:
    key = issue["key"]
    summary = issue["fields"].get("summary", key)
    log.info("=== Processing %s: %s ===", key, summary)

    attachments = jira.get_attachments(issue)
    translator   = jira.get_translator(issue)
    original_url = jira.get_original_url(issue)

    log.info("Translator: %s | Original URL: %s", translator or "(none)", original_url or "(none)")

    # ── 1. Find .docx translation ──────────────────────────────────────────────
    docx_files = [a for a in attachments if a["filename"].lower().endswith(".docx")]
    if not docx_files:
        log.warning("No .docx attachment in %s — skipping", key)
        return False
    docx_data = jira.download_attachment(docx_files[0])

    # ── 2. Extract Ukrainian title + body ──────────────────────────────────────
    uk_title, body_html = docx_to_html(docx_data)
    if not uk_title:
        uk_title = summary
    log.info("Extracted title: %s", uk_title)

    # ── 3. Fetch original article title ────────────────────────────────────────
    original_title = _fetch_page_title(original_url)
    log.info("Original title: %s", original_title or "(not found)")

    # ── 4. Generate tags ───────────────────────────────────────────────────────
    tags_list = generate_tags(uk_title, body_html)
    log.info("Tags: %s", tags_list)
    tag_ids = [wp.get_or_create_tag(t) for t in tags_list]

    # ── 5. Build full content with attribution ─────────────────────────────────
    attribution = build_attribution(original_url, original_title, translator)
    full_content = body_html + "\n" + attribution

    # ── 6. Upload featured image ───────────────────────────────────────────────
    featured_media_id = None
    img_result = pick_main_image(attachments, original_url, jira)
    if img_result:
        img_bytes, img_filename, img_mime = img_result
        featured_media_id = wp.upload_media(img_filename, img_bytes, img_mime)
        log.info("Uploaded image '%s' → media id %s", img_filename, featured_media_id)
    else:
        log.info("No image attachments found")

    # ── 7. Create WordPress draft post ────────────────────────────────────────
    wp_post = wp.create_post(
        title=uk_title,
        content=full_content,
        status="draft",
        featured_media=featured_media_id,
        tags=tag_ids,
    )
    wp_edit_url = (
        f"{WP_BASE_URL.rstrip('/')}/wp-admin/post.php"
        f"?post={wp_post['id']}&action=edit"
    )
    log.info("Created WP draft → %s", wp_edit_url)

    # ── 8. Comment on JIRA with WP link (idempotency aid) ─────────────────────
    jira.add_comment(key, f"WordPress draft created: {wp_edit_url}")

    # ── 9. Transition JIRA issue to 'In Review' ────────────────────────────────
    jira.transition_issue(key, JIRA_NEXT_STATUS)

    return True


# ── Entry point ────────────────────────────────────────────────────────────────

def main() -> None:
    jira = JiraClient()
    wp   = WpClient()

    issues = jira.get_issues_in_status(JIRA_TRIGGER_STATUS)
    log.info("Found %d issue(s) in '%s'", len(issues), JIRA_TRIGGER_STATUS)

    processed = 0
    for issue in issues:
        try:
            if process_issue(issue, jira, wp):
                processed += 1
        except Exception:
            log.exception("Failed to process %s", issue["key"])

    log.info("Done — processed %d / %d issue(s).", processed, len(issues))


if __name__ == "__main__":
    main()
