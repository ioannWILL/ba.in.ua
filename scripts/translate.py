#!/usr/bin/env python3
"""
Auto-translation workflow.

Finds the oldest JIRA issue in 'TO TRANSLATE' status (lowest BIU number),
translates its title and description from English to Ukrainian using Claude,
creates a .docx file, attaches it to the JIRA issue, and moves the issue
to 'TO REVIEW'.
"""

import logging
import os
import re

import anthropic
import requests
from docx import Document
from docx.shared import Pt

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
log = logging.getLogger(__name__)


# ── GitHub Actions job summary ─────────────────────────────────────────────────

def gha_summary(text: str) -> None:
    """Append markdown text to the GitHub Actions job summary (GITHUB_STEP_SUMMARY)."""
    path = os.environ.get("GITHUB_STEP_SUMMARY")
    if path:
        with open(path, "a") as f:
            f.write(text + "\n")

# ── Configuration ──────────────────────────────────────────────────────────────

JIRA_BASE_URL    = os.environ["JIRA_BASE_URL"]
JIRA_EMAIL       = os.environ["JIRA_EMAIL"]
JIRA_API_TOKEN   = os.environ["JIRA_API_TOKEN"]
JIRA_PROJECT_KEY = os.environ["JIRA_PROJECT_KEY"]

JIRA_SOURCE_STATUS  = os.environ.get("JIRA_TRANSLATE_SOURCE_STATUS")  or "TO TRANSLATE"
JIRA_WORKING_STATUS = os.environ.get("JIRA_TRANSLATE_WORKING_STATUS") or "IN TRANSLATION"
JIRA_DONE_STATUS    = os.environ.get("JIRA_TRANSLATE_DONE_STATUS")    or "TO REVIEW"

ANTHROPIC_API_KEY = os.environ["ANTHROPIC_API_KEY"]

_claude = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


# ── JIRA client ────────────────────────────────────────────────────────────────

class JiraClient:
    def __init__(self):
        self.session = requests.Session()
        self.session.auth = (JIRA_EMAIL, JIRA_API_TOKEN)
        self.session.headers.update({"Accept": "application/json"})
        self.base = JIRA_BASE_URL.rstrip("/")

    def _get(self, path, **kw):
        r = self.session.get(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r.json()

    def _post(self, path, **kw):
        r = self.session.post(f"{self.base}{path}", **kw)
        r.raise_for_status()
        return r

    def get_oldest_issue_in_status(self, status: str) -> dict | None:
        """Return the issue with the lowest BIU number in the given status."""
        jql = (
            f'project = "{JIRA_PROJECT_KEY}" AND status = "{status}" '
            f'ORDER BY key ASC'
        )
        r = self.session.post(
            f"{self.base}/rest/api/3/search/jql",
            json={
                "jql": jql,
                "maxResults": 1,
                "fields": ["summary", "description", "status", "assignee"],
            },
        )
        r.raise_for_status()
        issues = r.json().get("issues", [])
        return issues[0] if issues else None

    def get_transition_id(self, issue_key: str, status_name: str) -> str | None:
        data = self._get(f"/rest/api/3/issue/{issue_key}/transitions")
        for t in data.get("transitions", []):
            if t["to"]["name"].lower() == status_name.lower():
                return t["id"]
        return None

    def transition_issue(self, issue_key: str, status_name: str) -> bool:
        tid = self.get_transition_id(issue_key, status_name)
        if not tid:
            log.warning("Transition '%s' not found for %s", status_name, issue_key)
            return False
        self._post(
            f"/rest/api/3/issue/{issue_key}/transitions",
            json={"transition": {"id": tid}},
        )
        log.info("Transitioned %s → %s", issue_key, status_name)
        return True

    def attach_file(self, issue_key: str, filename: str, data: bytes, mime: str) -> None:
        """Upload a file as an attachment to a JIRA issue."""
        r = self.session.post(
            f"{self.base}/rest/api/3/issue/{issue_key}/attachments",
            files={"file": (filename, data, mime)},
            headers={"X-Atlassian-Token": "no-check"},
        )
        r.raise_for_status()
        log.info("Attached '%s' to %s", filename, issue_key)

    def add_comment(self, issue_key: str, text: str) -> None:
        self._post(
            f"/rest/api/3/issue/{issue_key}/comment",
            json={"body": {
                "type": "doc", "version": 1,
                "content": [{"type": "paragraph",
                             "content": [{"type": "text", "text": text}]}],
            }},
        )


# ── ADF → structured blocks ────────────────────────────────────────────────────

def _node_text(node: dict | list) -> str:
    """Extract plain text from an ADF node."""
    if isinstance(node, list):
        return "".join(_node_text(n) for n in node)
    if isinstance(node, dict):
        if node.get("type") == "text":
            return node.get("text", "")
        return "".join(_node_text(c) for c in node.get("content", []))
    return ""


def adf_to_blocks(node: dict | list) -> list[tuple[str, str]]:
    """Return list of (block_type, text) from an ADF document.

    block_type is one of: 'heading1'–'heading6', 'paragraph', 'listitem', 'blockquote'.
    Only non-empty blocks are returned.
    """
    blocks: list[tuple[str, str]] = []

    if isinstance(node, list):
        for n in node:
            blocks.extend(adf_to_blocks(n))
        return blocks

    if not isinstance(node, dict):
        return blocks

    t = node.get("type", "")

    if t in ("doc",):
        for child in node.get("content", []):
            blocks.extend(adf_to_blocks(child))

    elif t == "paragraph":
        text = _node_text(node).strip()
        if text:
            blocks.append(("paragraph", text))

    elif t == "heading":
        level = node.get("attrs", {}).get("level", 2)
        text = _node_text(node).strip()
        if text:
            blocks.append((f"heading{level}", text))

    elif t in ("bulletList", "orderedList"):
        for item in node.get("content", []):
            text = _node_text(item).strip()
            if text:
                blocks.append(("listitem", text))

    elif t == "blockquote":
        text = _node_text(node).strip()
        if text:
            blocks.append(("blockquote", text))

    elif t == "codeBlock":
        text = _node_text(node).strip()
        if text:
            blocks.append(("paragraph", text))  # keep as paragraph in docx

    return blocks


# ── Translation via Claude ─────────────────────────────────────────────────────

_SYSTEM_PROMPT = (
    "You are a professional translator. Translate the given English text to Ukrainian. "
    "Return ONLY the translated text — no explanations, no quotes, no additional commentary. "
    "Preserve the original meaning exactly. The text is from a business/analytics article "
    "and the author has granted permission for this translation."
)


def translate(text: str) -> str:
    """Translate a single text block English → Ukrainian using Claude."""
    if not text.strip():
        return text
    msg = _claude.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2048,
        system=_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": text}],
    )
    result = msg.content[0].text.strip()
    log.info("Translated: %s… → %s…", text[:60], result[:60])
    return result


# ── .docx builder ──────────────────────────────────────────────────────────────

def build_docx(title_uk: str, blocks_uk: list[tuple[str, str]]) -> bytes:
    """Create a .docx with the translated title and body blocks."""
    doc = Document()

    # Title as Heading 1
    heading = doc.add_heading(title_uk, level=1)
    heading.runs[0].font.size = Pt(16)

    for block_type, text in blocks_uk:
        if block_type.startswith("heading"):
            level = int(block_type[-1])
            doc.add_heading(text, level=min(level, 9))
        elif block_type == "listitem":
            para = doc.add_paragraph(text, style="List Bullet")
        elif block_type == "blockquote":
            para = doc.add_paragraph(text)
            para.paragraph_format.left_indent = Pt(24)
        else:
            doc.add_paragraph(text)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    jira = JiraClient()

    gha_summary("## 🔤 Auto-translation\n")

    # 1. Find oldest issue in TO TRANSLATE
    issue = jira.get_oldest_issue_in_status(JIRA_SOURCE_STATUS)
    if not issue:
        msg = f"No issues found in **{JIRA_SOURCE_STATUS}**. Nothing to translate."
        log.info(msg)
        gha_summary(f"ℹ️ {msg}")
        return

    key     = issue["key"]
    fields  = issue["fields"]
    summary = fields.get("summary", key)
    issue_number = re.search(r"\d+", key).group()
    jira_url = f"{JIRA_BASE_URL.rstrip('/')}/browse/{key}"

    log.info("Found issue: %s — %s", key, summary)
    gha_summary(f"| | |\n|---|---|\n"
                f"| **Issue** | [{key}]({jira_url}) |\n"
                f"| **Summary** | {summary} |\n")

    # 2. Move to IN TRANSLATION
    ok = jira.transition_issue(key, JIRA_WORKING_STATUS)
    gha_summary(f"| **Step 1 — Move to {JIRA_WORKING_STATUS}** | {'✅' if ok else '⚠️ transition not found'} |")

    # 3. Translate title
    log.info("Translating title…")
    title_uk = translate(summary)
    gha_summary(f"| **Step 2 — Title (UK)** | {title_uk} |")

    # 4. Extract and translate description blocks
    desc = fields.get("description")
    blocks_en: list[tuple[str, str]] = []
    if isinstance(desc, dict):
        blocks_en = adf_to_blocks(desc)
    elif isinstance(desc, str) and desc.strip():
        blocks_en = [("paragraph", p.strip()) for p in desc.split("\n\n") if p.strip()]

    log.info("Translating %d block(s) from description…", len(blocks_en))
    gha_summary(f"| **Step 3 — Description blocks** | {len(blocks_en)} block(s) found |")

    if not blocks_en:
        log.warning("Description is empty — .docx will contain only the title.")
        gha_summary("| | ⚠️ Description is empty — only title will be in the .docx |")

    blocks_uk = [(btype, translate(text)) for btype, text in blocks_en]

    # 5. Build .docx
    docx_bytes = build_docx(title_uk, blocks_uk)
    filename = f"translation_biu{issue_number}.docx"
    log.info("Created %s (%d bytes)", filename, len(docx_bytes))
    gha_summary(f"| **Step 4 — .docx created** | `{filename}` ({len(docx_bytes):,} bytes) |")

    # 6. Attach .docx to JIRA issue
    jira.attach_file(
        key, filename, docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    jira.add_comment(key, f"Translation file attached: {filename}")
    gha_summary(f"| **Step 5 — Attached to JIRA** | ✅ [{key}]({jira_url}) |")

    # 7. Move to TO REVIEW
    ok = jira.transition_issue(key, JIRA_DONE_STATUS)
    gha_summary(f"| **Step 6 — Move to {JIRA_DONE_STATUS}** | {'✅' if ok else '⚠️ transition not found'} |")

    log.info("Done — %s translated and moved to '%s'.", key, JIRA_DONE_STATUS)
    gha_summary(f"\n**✅ Done — [{key}]({jira_url}) moved to {JIRA_DONE_STATUS}**")


if __name__ == "__main__":
    main()
